from rest_framework import viewsets, permissions, mixins
from rest_framework.response import Response
from rest_framework.decorators import action
from .models import MeetingRoom, Reservation
from .serializers import MeetingRoomSerializer, ReservationSerializer
from django.utils import timezone
from django.http import HttpResponse
from django.utils.dateparse import parse_date

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class MeetingRoomViewSet(
    viewsets.GenericViewSet,
    mixins.ListModelMixin,
    mixins.CreateModelMixin,
    mixins.UpdateModelMixin,
    mixins.DestroyModelMixin,
):
    queryset = MeetingRoom.objects.all()
    serializer_class = MeetingRoomSerializer

    def create(self, request, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        self.perform_create(serializer)
        return Response({'id': serializer.instance.id}, status=201)


class ReservationViewSet(
    viewsets.GenericViewSet,
    mixins.ListModelMixin,
    mixins.CreateModelMixin,
    mixins.DestroyModelMixin,
):
    queryset = Reservation.objects.all()
    serializer_class = ReservationSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        return Reservation.objects.filter(start_time__date=timezone.now().date())

    def perform_create(self, serializer):
        serializer.save()

    # GET reservations/availability/?room_id=1
    # GET reservations/availability/?room_id=1&date=2024-09-10
    @action(detail=False, methods=['get'])
    def availability(self, request):
        """
            Получить все записи о бронировании переговорной
            комнаты на текущий день (либо на определенное время).
        """

        date_str = request.query_params.get('date', timezone.now().date())
        room_id = request.query_params.get('room_id')
        start_time = request.query_params.get('start_time')
        end_time = request.query_params.get('end_time')
        available_param = request.query_params.get('available')

        if isinstance(date_str, str):
            date = parse_date(date_str)
        else:
            date = timezone.now().date()

        if start_time and end_time:
            start_time = timezone.datetime.fromisoformat(start_time)
            end_time = timezone.datetime.fromisoformat(end_time)
        else:
            start_time = timezone.datetime.combine(date, timezone.datetime.min.time())
            end_time = timezone.datetime.combine(date, timezone.datetime.max.time())

        # Получаем все бронирования для указанной переговорной комнаты за указанный день
        reservations = Reservation.objects.filter(room__id=room_id, start_time__date=date)

        if available_param in ['false', 'False', '0']:
            available = not reservations.exists()
        else:
            available = reservations.exists()

        reservation_details = []
        if reservations.exists():
            reservation_details = [{
                    'user': res.user.username,
                    'start_time': res.start_time,
                    'end_time': res.end_time
                } for res in reservations
            ]

        return Response({
            'room_id': room_id,
            'date': date,
            'available': available,
            'reservations': reservation_details,
        })

    # reservations/report/?room_number=1&start_date=2024-09-01&end_date=2024-09-10
    @action(detail=False, methods=['get'])
    def report(self, request):
        """
            Получает отчет в формате word, содержащий в себе данные за определенный
            период о бронированиях переговорных комнат (или определенной комнаты).
            В отчете содержаться данные о том, кто бронировал, в какое время
            и для каких целей.
        """

        room_number = request.query_params.get('room_number')
        start_date = request.query_params.get('start_date')
        end_date = request.query_params.get('end_date')

        if not room_number or not start_date or not end_date:
            return Response({'error': 'Введите комнату, время начала, и время конца.'}, status=400)
        
        reservations = Reservation.objects.filter(
            room__id=room_number,
            start_time__date__range=[start_date, end_date]
        )

        report = Document()
        report.add_heading('Бронирование Переговорных Комнат', 0)

        if reservations.exists():
            report.add_heading(f'Отчет по комнате {reservations[0].room.room_number}', level=1)
            report.add_paragraph().add_run('\n')

            for reservation in reservations:
                paragraph = report.add_paragraph()
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                user_run = paragraph.add_run(f"Бронирование от: {reservation.user.username}\n")
                user_run.bold = True
                
                paragraph.add_run(
                    f'С {reservation.start_time.strftime("%d-%m-%Y %H:%M")} до {reservation.end_time.strftime("%d-%m-%Y %H:%M")}\n'
                )
                paragraph.add_run(
                    f'Цель бронирования: {reservation.purpose}\n'
                )
                paragraph.add_run('-' * 70 + '\n')
                
            footer = report.add_paragraph()
            footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            footer.add_run(f'Сгенерировано {timezone.now().strftime("%d-%m-%Y %H:%M:%S")}')
        else:
            report.add_paragraph('Нет доступных данных за указанный период.', style='Normal')

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=report.docx'
        report.save(response)
        
        return response
